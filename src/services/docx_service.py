import re
import io
import logging
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import RGBColor

from src.models.langg import MemoRequest

logger = logging.getLogger(__name__)

PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")
MEMO_DIR = Path("memo")
TEMPLATE_PATH = MEMO_DIR / "memo_template.docx"


class DOCXService:

    def __init__(
        self,
        template_path: Path = TEMPLATE_PATH,
        output_dir: Path = MEMO_DIR,
    ):
        self.template_path = template_path
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ── Public ────────────────────────────────────────────────────────────────

    def generate(self, memo: MemoRequest) -> Path:
        flat = self._flatten(memo.model_dump())
        doc  = self._fill_template(flat)
        path = self._build_output_path(memo.investment_memo.project_name)
        doc.save(path)
        logger.info(f"IC Memo guardado en: {path}")
        return path

    def generate_bytes(self, memo: MemoRequest) -> bytes:
        flat = self._flatten(memo.model_dump())
        doc  = self._fill_template(flat)
        buf  = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ── Private ───────────────────────────────────────────────────────────────

    def _flatten(self, obj: dict | list, prefix: str = "") -> dict[str, str]:
        result: dict[str, str] = {}

        if isinstance(obj, dict):
            for k, v in obj.items():
                key = f"{prefix}.{k}" if prefix else k
                if isinstance(v, (dict, list)):
                    result.update(self._flatten(v, key))
                elif v is not None:
                    result[key] = str(v)

        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                key = f"{prefix}.{i}" if prefix else str(i)
                if isinstance(item, (dict, list)):
                    result.update(self._flatten(item, key))
                elif item is not None:
                    result[key] = str(item)

        return result

    def _resolve(self, key: str, flat: dict[str, str]) -> str:
        val = flat.get(key.strip())
        if val is None:
            logger.warning(f"Placeholder did not resolve: '{key.strip()}'")
        return val if val is not None else ""

    def _merge_runs(self, paragraph) -> None:
        """
        Word fragmenta los placeholders en múltiples runs consecutivos,
        por ejemplo: run1='{{budget', run2='.land', run3='}}'.
        Esto hace que el regex nunca haga match.

        Consolida el texto de todos los runs en el primero y vacía el resto,
        preservando el formato del primer run. Solo actúa cuando el texto
        concatenado contiene al menos un placeholder.
        """
        runs = paragraph.runs
        if len(runs) < 2:
            return

        full_text = "".join(r.text for r in runs)
        if "{{" not in full_text:
            return

        runs[0].text = full_text
        for run in runs[1:]:
            run.text = ""

    def _replace_runs(self, paragraphs, flat: dict[str, str]) -> None:
        for para in paragraphs:
            self._merge_runs(para)

            for run in para.runs:
                if "{{" not in run.text:
                    continue
                new_text = PLACEHOLDER_RE.sub(
                    lambda m: self._resolve(m.group(1), flat),
                    run.text
                )
                if new_text != run.text:
                    run.text = new_text
                    if run.font.color.type is not None:
                        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    def _fill_template(self, flat: dict[str, str]) -> Document:
        if not self.template_path.exists():
            raise FileNotFoundError(
                f"Template not found: {self.template_path}. "
                "Make sure memo_template.docx is in the memo directory."
            )

        doc = Document(self.template_path)

        def process_table(table):
            for row in table.rows:
                for cell in row.cells:
                    self._replace_runs(cell.paragraphs, flat)
                    for nested in cell.tables:
                        process_table(nested)

        self._replace_runs(doc.paragraphs, flat)
        for table in doc.tables:
            process_table(table)

        return doc

    def _build_output_path(self, project_name: str) -> Path:
        safe_name = re.sub(r"[^\w\-]", "_", project_name)[:60]
        timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename   = f"IC_Memo_{safe_name}_{timestamp}.docx"
        return self.output_dir / filename